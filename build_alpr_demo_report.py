from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r"D:\YOLO_ALPR_Project")
OUTPUT = PROJECT / "ALPR_技术演示与阶段汇报.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "1F6B45"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_run_font(run, name: str = "Microsoft YaHei", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_paragraph_border(paragraph, bottom: str = "D7DBE2") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), bottom)
    borders.append(border)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_text(doc: Document, text: str, style: str | None = None, color: str | None = None, bold: bool = False, italic: bool = False, align=None) -> object:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=11, color=color or "1D2530", bold=bold, italic=italic)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="1D2530")


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="1D2530")


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.keep_together = True
    set_paragraph_border(p, "D7DBE2")
    r = p.add_run(code)
    set_run_font(r, name="Consolas", size=8.6, color="18212B")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F9")
    p._p.get_or_add_pPr().append(shading)


def add_callout(doc: Document, title: str, body: str, tone: str = "blue") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    fill = {"blue": CALLOUT, "green": "EDF7F1", "gold": "FFF8E8", "red": "FCEEEE"}[tone]
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color={"blue": DARK_BLUE, "green": GREEN, "gold": GOLD, "red": RED}[tone], bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10.2, color="1D2530")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.1) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.4, color="1D2530")
            if index in (0, len(row) - 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1D2530")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Code" not in styles:
        styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("YOLO ALPR 项目 | 技术演示与阶段汇报")
    set_run_font(r, size=9, color=MUTED)
    set_paragraph_border(header, "D7DBE2")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("技术汇报与现场演示手册")
    set_run_font(r, size=13, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("YOLO ALPR 中文车牌识别系统")
    set_run_font(r, size=28, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Windows 视频验证、RK3588 部署、MIPI 摄像头、OBB 定位与车牌恢复模型")
    set_run_font(r, size=14, color=MUTED)
    p.paragraph_format.space_after = Pt(30)
    add_callout(
        doc,
        "本阶段目标",
        "在真实高速道路场景中，从视频或 MIPI 摄像头稳定提取车牌候选、进行中文 OCR 投票，并为后续 OBB 微调与去模糊恢复保留可复核的证据链。",
        "blue",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目根目录：D:\\YOLO_ALPR_Project\n汇报日期：2026-06-21")
    set_run_font(r, size=10.5, color=MUTED)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("目录与阅读方式", level=1)
    add_number(doc, "系统目标与当前结论")
    add_number(doc, "总体架构与端到端流程")
    add_number(doc, "Windows 端视频识别：Top-K、OCR 与报告")
    add_number(doc, "RK3588 板端视频与 MIPI 摄像头识别")
    add_number(doc, "车牌 OBB 微调：数据、训练、A/B 结论")
    add_number(doc, "车牌去模糊模型：数据、训练与静态验证")
    add_number(doc, "困难、排查过程与工程决策")
    add_number(doc, "现场演示 Runbook 与命令清单")
    add_number(doc, "后续路线与风险控制")
    add_callout(doc, "汇报建议", "先演示 Windows Top-K 报告，再演示板端视频的锁定与预测，最后接入 MIPI 实时画面。OBB 微调和去模糊部分以“已完成实验 + 保守结论”方式说明，避免把验证集指标误当成实际收益。", "gold")


def add_overview(doc: Document) -> None:
    doc.add_heading("1. 系统目标与当前结论", level=1)
    add_text(doc, "目标不是仅在论文数据集上获得较好指标，而是在 D:\\YOLO_ALPR_Project\\测试图\\14.mp4 所代表的真实道路场景中，稳定识别远距离、角度变化、运动模糊和中文车牌。系统同时需要具备 Windows 验证能力与 RK3588 板端部署能力。")
    add_table(
        doc,
        ["模块", "当前状态", "阶段结论"],
        [
            ["Windows Top-K", "可运行", "为每辆车保存 Top-K 原始帧、车辆图、透视车牌图和 JSON 评分。"],
            ["OCR 与投票", "可运行", "HyperLPR3 用于候选 OCR；按字符加权投票后锁定。"],
            ["RK3588 视频", "可运行", "RKNN 车辆/OBB + HyperLPR3；支持锁定、预测显示和结果保存。"],
            ["RK3588 MIPI", "可运行", "使用原生 v4l2-ctl mmap 采集；空场景实测约 18.17 FPS。"],
            ["OBB e20", "已训练/未替换", "25 张历史候选 A/B 未优于基线 best_obb.pt，继续使用基线。"],
            ["去模糊模型", "已训练/暂不接入", "配对指标提升，但真实模糊车牌 OCR 未稳定提升，保留为后续 Top-K 离线增强。"],
        ],
        [1500, 2100, 5760],
    )
    add_callout(doc, "关键结论", "本阶段已打通“检测 - 车牌候选 - OCR - 投票 - 锁定 - 板端显示 - 可复核输出”的完整闭环。对于效果未明确改善的 OBB 微调和去模糊模型，采用保守策略：不替换线上基线，保留数据和训练管线继续迭代。", "green")


def add_architecture(doc: Document) -> None:
    doc.add_heading("2. 总体架构与端到端流程", level=1)
    add_text(doc, "系统以车辆检测为入口，以旋转框 OBB 定位车牌，以透视拉正后的 320x96 车牌图作为 OCR、质量评分、去模糊和后续 RKNN 量化的共同接口。")
    add_table(
        doc,
        ["阶段", "输入", "处理", "输出"],
        [
            ["车辆检测", "视频帧 / MIPI 帧", "YOLO/RKNN 检测 car、bus、truck；IoU 跟踪", "track_id、车辆框、置信度"],
            ["进入门控", "车辆框", "按车辆置信度、宽高、面积、面积比判断是否值得跑车牌 OBB", "减少远车无效推理"],
            ["OBB 定位", "车辆 ROI", "旋转框车牌检测、几何过滤、透视拉正", "320x96 plate crop"],
            ["候选排序", "plate crop", "面积、清晰度、曝光、对比度、OBB 置信度加权", "每车 Top-K"],
            ["OCR 与投票", "车辆图/车牌图", "HyperLPR3、字符位置加权投票、锁定阈值", "稳定车牌文本"],
            ["板端展示", "轨迹与锁定文本", "绿框 LOCK、橙框 PRED、中文渲染", "实时预览和标注视频"],
        ],
        [1500, 1900, 3300, 2660],
    )
    add_text(doc, "Top-K 质量分默认权重：车牌面积 0.35、Laplacian 清晰度 0.30、曝光 0.15、对比度 0.10、OBB 置信度 0.10。OCR 不作为第一版候选质量分，以避免 OCR 推理成为候选抽取的干扰因素。", italic=True, color=MUTED)


def add_windows(doc: Document) -> None:
    doc.add_heading("3. Windows 端视频识别与候选报告", level=1)
    doc.add_heading("3.1 为什么采用 Top-K 候选", level=2)
    add_text(doc, "单帧 fallback 车牌经常恰好落在运动模糊或错误 OBB 时刻。Windows 端不急于对第一帧下结论，而是按 track_id 跨帧收集候选，保留最清晰、最大、曝光更合理的 Top-K，再进行 OCR 复核和人工报告浏览。")
    add_bullet(doc, "独立脚本 alpr_topk_capture.py，不改动原 alpr_sahi_snapshot2.py 链路。")
    add_bullet(doc, "支持 IoU tracker；Windows 端可选 Ultralytics ByteTrack，但当前验证核心不依赖 tracker 额外安装。")
    add_bullet(doc, "支持 live OCR、字符加权投票、锁定后短时预测跟随和 rejected 原因归档。")
    doc.add_heading("3.2 演示命令：Windows 实时窗口", level=2)
    add_code(doc, "cd D:\\YOLO_ALPR_Project\nD:\\miniconda\\envs\\alpr_env\\python.exe alpr_topk_capture.py `\n  --video D:\\YOLO_ALPR_Project\\测试图\\14.mp4 `\n  --plate-model D:\\YOLO_ALPR_Project\\best_obb.pt `\n  --with-ocr --live-ocr --ocr-engine hyperlpr3 `\n  --show-window --show-waiting --progress-interval 300")
    add_text(doc, "操作说明：窗口中 q 可提前结束；Ctrl+C 也会进入保护逻辑，保存当前 Top-K 与 summary.json。锁定后绿色文本为最终投票结果，橙色文本为短时预测状态。")
    doc.add_heading("3.3 演示命令：生成 Top-K HTML 报告", level=2)
    add_code(doc, "D:\\miniconda\\envs\\alpr_env\\python.exe topk_report.py `\n  --run-dir D:\\YOLO_ALPR_Project\\captures_topk\\run_YYYYMMDD_HHMMSS `\n  --with-ocr --ocr-engine hyperlpr3 --overwrite-ocr")
    add_text(doc, "报告会生成 topk_report.html、topk_report.csv 和 best_ocr_summary.txt。浏览重点：车牌 crop 是否可读、OBB 是否落在真实车牌、OCR/raw OCR 的差异、质量分及 rejected 原因。")
    add_figure(doc, PROJECT / "captures_topk" / "run_20260617_211958" / "track_10" / "plate_rank1.jpg", "图 1. Windows Top-K 车牌候选示例（track_10）", 4.0)


def add_rk3588(doc: Document) -> None:
    doc.add_heading("4. RK3588 板端视频与 MIPI 摄像头识别", level=1)
    doc.add_heading("4.1 板端部署组成", level=2)
    add_table(
        doc,
        ["项目", "板端位置 / 状态", "说明"],
        [
            ["设备", "192.168.137.168 / root", "RK3588，Debian 11，RKNPU v2。"],
            ["运行目录", "/root/alpr_topk_rk3588", "Python 脚本、models 软链接、runs 输出。"],
            ["模型", "vehicle.rknn + best_obb.rknn", "车辆检测与 OBB 车牌定位。"],
            ["OCR", "HyperLPR3 + ONNX Runtime", "板端已验证可运行；用于候选/车辆图 OCR。"],
            ["中文显示", "Pillow 11.3.0 + 文泉驿正黑", "锁定与预测标签可显示汉字。"],
        ],
        [1500, 3300, 4560],
    )
    doc.add_heading("4.2 演示命令：板端视频文件", level=2)
    add_code(doc, "ssh -i \"C:\\Users\\Lenovo\\.ssh\\rk3588_alpr\" root@192.168.137.168\ncd /root/alpr_topk_rk3588\npython3 rk3588_topk_capture.py `\n  --video /root/deploy/144.mp4 `\n  --vehicle-model models/vehicle.rknn `\n  --plate-model models/best_obb.rknn --plate-imgsz 640 `\n  --ocr-engine hyperlpr3 --output runs --save-video")
    add_text(doc, "144.mp4 与 Windows 端 14.mp4 为同一演示场景。运行结束后，runs/run_时间戳/annotated.mp4 为标注视频；各 track 目录仅在有有效候选时才会创建。")
    doc.add_heading("4.3 锁定、预测与显示验证", level=2)
    add_text(doc, "验证中临时将 vote-threshold 设为 1，只为快速触发 LOCK 并检验中文显示；实际演示请使用默认 3 次有效 OCR 投票，避免单帧误识别直接锁定。")
    add_figure(doc, PROJECT / "RK3588_dev" / "lock_label_probe" / "lock_60.jpg", "图 2. 板端视频：绿色 LOCK 中文车牌标签与车辆框", 6.1)
    add_figure(doc, PROJECT / "RK3588_dev" / "lock_label_probe" / "lock_55.jpg", "图 3. 板端视频：橙色 PRED 中文标签，表示检测间隔帧的短时运动预测", 6.1)
    doc.add_heading("4.4 演示命令：MIPI 摄像头 + 浏览器预览", level=2)
    add_text(doc, "终端 A 启动浏览器 MJPEG 服务：")
    add_code(doc, "cd /root/deploy\npython3 stream.py")
    add_text(doc, "终端 B 启动实时识别：")
    add_code(doc, "cd /root/alpr_topk_rk3588\npython3 rk3588_topk_capture.py `\n  --video mipi --camera-device /dev/video22 `\n  --camera-width 1920 --camera-height 1080 `\n  --vehicle-model models/vehicle.rknn `\n  --plate-model models/best_obb.rknn --plate-imgsz 640 `\n  --ocr-engine hyperlpr3 --vehicle-detect-interval 2 `\n  --output runs --publish-frame /tmp/frame.jpg --publish-interval 2")
    add_text(doc, "Windows 浏览器打开：http://192.168.137.168:8080。对于车速较快的近距离车辆使用 --vehicle-detect-interval 2；默认 3 更省算力。Ctrl+C 会保存当前结果。")
    add_callout(doc, "MIPI 色彩说明", "video22 实际输出为 UYVY。脚本已强制请求并显式解码 UYVY；原始帧仍偏绿，说明问题位于 rkaiq_3A_server/ISP 白平衡或相机 IQ 文件，而不是网页预览或 BGR/RGB 通道错置。可临时追加 --mipi-gray-world 对比软件补偿，但不建议未现场复核就默认开启。", "gold")


def add_obb(doc: Document) -> None:
    doc.add_heading("5. OBB 微调：数据构建、训练与 A/B 结论", level=1)
    add_text(doc, "为改善护栏侧、远距离和倾斜车牌的定位，构建了统一 YOLO-OBB 单类别数据集。CCPD 四边形由文件名读取，CRPD single/double/multi 的四边形由标签文本读取；真实视频 hard case 可经人工 review 再合并。")
    add_table(
        doc,
        ["数据/实验", "内容", "用途"],
        [
            ["plate_dataset_obb_finetune", "训练 30,786 张；验证 7,255 张；约 17.6 GB", "Colab OBB 微调。"],
            ["CCPD2020", "基础中文车牌场景", "覆盖常规视角与车牌形式。"],
            ["CRPD single/double/multi", "真实复杂道路与多车牌条件", "增强复杂定位鲁棒性。"],
            ["视频 hard case review", "从 Top-K/rejected 导出，四点人工标注", "针对真实失败样本补强。"],
        ],
        [2300, 3900, 3160],
    )
    doc.add_heading("5.1 Colab 训练命令", level=2)
    add_code(doc, "# Colab GPU Runtime\n!pip -q install -U ultralytics\n!cp \"$DRIVE_ROOT/train_obb_colab.py\" /content/train_obb_colab.py\n!unzip -q \"$DRIVE_ROOT/plate_dataset_obb_finetune.zip\" -d /content\n!python /content/train_obb_colab.py")
    add_text(doc, "首轮配置：imgsz=640、epochs=20、batch=16、AdamW。训练第 8 轮已经达到较高验证指标，最终获得 best_obb_finetuned_e20.pt。")
    doc.add_heading("5.2 实际视频 A/B，而不是只看 mAP", level=2)
    add_table(
        doc,
        ["比较项", "best_obb.pt", "best_obb_finetuned_e20.pt", "结论"],
        [
            ["保存的历史候选帧", "25 张中检出 23 张", "25 张中检出 23 张", "无新增检出。"],
            ["共同检出置信度均值", "0.938", "0.741", "e20 在 23/23 样本更低。"],
            ["track_10 / track_15", "定位正确", "框基本重合", "未显示明显改善。"],
            ["track_12", "JPEG A/B 阈值 0.25 下未检出", "同样未检出", "未解决关键错误定位问题。"],
        ],
        [2250, 2200, 2600, 2310],
    )
    add_callout(doc, "模型选择决策", "尽管 e20 的验证集指标较高，真实 14.mp4 历史候选 A/B 没有改善，因此没有替换板端 best_obb.pt。这是刻意的工程保守策略：线上模型必须以目标场景的定位质量和 OCR 后果为准。", "red")
    add_figure(doc, PROJECT / "obb_ab_single_frames" / "run_20260621_093940" / "track_10_frame_3186_rank_1" / "finetuned_e20_overlay.jpg", "图 4. OBB 单帧 A/B 示例：e20 与基线框位置接近，未形成可替换优势", 6.1)


def add_deblur(doc: Document) -> None:
    doc.add_heading("6. 车牌去模糊模型：数据、训练与静态验证", level=1)
    add_text(doc, "去模糊不局限于 TinyUNet。针对 RK3588 后续量化部署，构建了固定 320x96 的轻量 PlateRestoreNet-Lite：残差式小型 U-Net，不使用扩散模型，目标是增强已有笔画并尽量避免“看似清晰但改写字符”的风险。")
    add_table(
        doc,
        ["项目", "内容", "结论"],
        [
            ["训练数据", "dataset + MDLP_Mini 组成配对数据", "train 9,340；val 1,027；test 1,041。"],
            ["模型", "约 203k 参数；输入/输出固定 1x3x96x320 BGR", "适合后续 ONNX/RKNN INT8 方向。"],
            ["测试集指标", "原始 PSNR 20.44 / SSIM 0.634；恢复后 PSNR 24.05 / SSIM 0.803", "配对测试提升明确。"],
            ["真实视频 crop", "部分 OCR 不升反降，错误 OBB 不能靠恢复修复", "暂不进入实时主链路。"],
        ],
        [1600, 4560, 3200],
    )
    doc.add_heading("6.1 静态图片演示命令", level=2)
    add_code(doc, "cd D:\\YOLO_ALPR_Project\nD:\\miniconda\\envs\\alpr_env\\python.exe test_deblur_image.py `\n  --input \"D:\\YOLO_ALPR_Project\\captures\\1_JC52IG_fallback_plate.jpg\" `\n  --model \"D:\\YOLO_ALPR_Project\\blur models\\plate_restore_lite_320x96.onnx\" `\n  --output \"D:\\YOLO_ALPR_Project\\blur models\\demo_output\"")
    add_text(doc, "提示：导出的 ONNX 可能带有同目录的 plate_restore_lite_320x96.onnx.data 外部权重文件，演示时两者必须同时存在。也可将 --model 指向 plate_restore_lite_best.pt 进行 PyTorch checkpoint 静态测试。")
    add_figure(doc, PROJECT / "blur models" / "onnx_static_test" / "1_JC52IG_fallback_plate_comparison.jpg", "图 5. 去模糊模型静态对比：左为输入，右为恢复结果。视觉增强不等于 OCR 一定更正确。", 5.4)
    add_callout(doc, "接入原则", "去模糊模型下一阶段只应对每辆车保存的 Top-K 车牌 crop 进行离线或准离线复核，并记录 HyperLPR3 恢复前后结果。只有真实视频 OCR 有稳定净提升时，才接入实时板端路径。", "gold")


def add_difficulties(doc: Document) -> None:
    doc.add_heading("7. 关键困难、尝试与工程决策", level=1)
    add_table(
        doc,
        ["问题", "已做尝试", "当前结论 / 决策"],
        [
            ["论文数据集泛化有限", "TinyUNet 只在特定论文数据集上效果较好；引入 CCPD、CRPD、MDLP。", "数据与真实道路场景差异是核心，需以 14.mp4 A/B 决策。"],
            ["单帧车牌太糊", "Top-K、质量评分、OCR 复核、投票。", "优先选择视频中最有价值帧，而非立即做生成式恢复。"],
            ["OCR 乱码 / 泛中文", "Paddle/plate-rec/ONNXRuntime 对比，最终使用 HyperLPR3。", "HyperLPR3 对当前车牌候选更实用；多帧投票降低偶发误识。"],
            ["OBB 错框", "车牌面积、宽高比、车辆重叠过滤；训练 e20；静态 A/B。", "e20 未改善目标样本，保留 baseline，准备 hard-case 人工标注。"],
            ["锁定后仍慢", "锁定跳过 OBB/OCR；新增车辆检测降频与预测。", "此前仍每帧跑 vehicle RKNN；现在间帧预测。"],
            ["板端中文不显示", "Pillow + 文泉驿字体渲染缓存。", "已实测 LOCK/PRED 中文标签可显示。"],
            ["空 track 目录", "保存阶段区分 tracks 与 ignored_tracks。", "无有效候选不再建目录。"],
            ["MIPI FPS 低", "OpenCV V4L2、原生 v4l2-ctl、GStreamer 探测。", "GStreamer 未编译；v4l2-ctl mmap 将脚本空场景提升至 18.17 FPS。"],
            ["MIPI 色偏", "自动转换与显式 UYVY 对比、灰度世界补偿、AIQ 服务检查。", "色偏来自 ISP/3A/IQ；脚本提供可控补偿，需现场调参。"],
        ],
        [2050, 3700, 3610],
    )
    add_text(doc, "板端还有一项部署风险：当前 .rknn 模型由 RKNN Toolkit 2.3.2 生成，而板端 librknnrt 为 1.6.0。运行时出现版本不匹配警告，当前仍可推理，但后续应使用与 Runtime 1.6 对应的 Toolkit 重新导出，避免性能和兼容性不确定性。", color=RED, bold=True)


def add_runbook(doc: Document) -> None:
    doc.add_heading("8. 现场演示 Runbook", level=1)
    add_callout(doc, "演示顺序", "推荐 12-15 分钟：1) Windows Top-K；2) HTML 报告；3) 板端 144.mp4；4) MIPI 浏览器预览；5) OBB A/B 结论；6) 去模糊静态对比与保守接入策略。", "blue")
    doc.add_heading("8.1 演示前检查", level=2)
    add_bullet(doc, "Windows 激活 alpr_env，并确认 best_obb.pt、yolo11n.pt、HyperLPR3、ONNX Runtime 可用。")
    add_bullet(doc, "确认 RK3588 可 SSH：ssh -i \"C:\\Users\\Lenovo\\.ssh\\rk3588_alpr\" root@192.168.137.168。")
    add_bullet(doc, "确认板端 /root/alpr_topk_rk3588/models 下的 vehicle.rknn 与 best_obb.rknn 软链接可读取。")
    add_bullet(doc, "确认 MIPI 设备存在：v4l2-ctl -d /dev/video22 --get-fmt-video，应显示 1920x1080 UYVY。")
    doc.add_heading("8.2 最短可复现命令", level=2)
    add_code(doc, "# Windows：完整视频候选与实时窗口\npython alpr_topk_capture.py --video D:\\YOLO_ALPR_Project\\测试图\\14.mp4 --with-ocr --live-ocr --ocr-engine hyperlpr3 --show-window")
    add_code(doc, "# Windows：对已完成 run 生成报告\npython topk_report.py --run-dir D:\\YOLO_ALPR_Project\\captures_topk\\run_YYYYMMDD_HHMMSS --with-ocr --ocr-engine hyperlpr3 --overwrite-ocr")
    add_code(doc, "# RK3588：视频文件\ncd /root/alpr_topk_rk3588 && python3 rk3588_topk_capture.py --video /root/deploy/144.mp4 --vehicle-model models/vehicle.rknn --plate-model models/best_obb.rknn --ocr-engine hyperlpr3 --output runs --save-video")
    add_code(doc, "# RK3588：MIPI 实时识别\ncd /root/alpr_topk_rk3588 && python3 rk3588_topk_capture.py --video mipi --camera-device /dev/video22 --camera-width 1920 --camera-height 1080 --vehicle-model models/vehicle.rknn --plate-model models/best_obb.rknn --ocr-engine hyperlpr3 --vehicle-detect-interval 2 --output runs --publish-frame /tmp/frame.jpg --publish-interval 2")
    doc.add_heading("8.3 演示时要说清的边界", level=2)
    add_bullet(doc, "单次 OCR 锁定仅用于渲染验证；正式结果依赖默认多帧投票。")
    add_bullet(doc, "去模糊模型在合成/配对测试指标上提升，但尚未证明真实视频 OCR 稳定受益。")
    add_bullet(doc, "OBB e20 的训练指标较好，但真实目标视频 A/B 没有收益，因此未替换基线。")
    add_bullet(doc, "MIPI 采集链路已优化到约 18 FPS 空场景；达到更高稳定帧率需继续做 RGA/原生采集或 C++ 路径优化。")


def add_roadmap(doc: Document) -> None:
    doc.add_heading("9. 后续路线与风险控制", level=1)
    add_table(
        doc,
        ["优先级", "下一步", "验收标准"],
        [
            ["P0", "用 rkaiq 现场校准白平衡/IQ；在道路光照下记录原图与灰度世界对比。", "不再出现绿/紫大幅偏色，车牌蓝色与白车颜色自然。"],
            ["P0", "以默认 3 次投票跑 MIPI 实车，记录锁定文本、预测框、漏检。", "锁定后框跟随稳定；错误锁定可复盘 vote_history。"],
            ["P1", "导出 track_10/12/15 等 hard case，人工标注真实四点 OBB，再训练第二阶段。", "新 OBB 在同一 14.mp4 A/B 中实际减少错框/漏检。"],
            ["P1", "将当前 .rknn 用 Toolkit 1.6 重新转换，匹配板端 Runtime 1.6。", "消除版本警告，验证输出一致且性能无回退。"],
            ["P2", "去模糊仅处理 Top-K，对 HyperLPR3 前后结果做车辆级统计。", "真实场景净提升后再考虑板端接入。"],
            ["P2", "评估 C++/RGA 或更直接 V4L2 mmap 采集，减少 Python pipe 拷贝。", "实时预览、车辆检测和锁定跟随进一步接近 25 FPS。"],
        ],
        [1000, 5000, 3360],
    )
    add_callout(doc, "最终原则", "所有模型与参数替换都必须以目标道路视频的候选质量、正确 OCR、错误锁定率和板端稳定性为验收依据。验证集 mAP、PSNR、SSIM 是必要参考，但不是部署决策的充分条件。", "green")


def build() -> None:
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    doc.core_properties.title = "YOLO ALPR 技术演示与阶段汇报"
    doc.core_properties.subject = "Windows、RK3588、MIPI、OBB 微调与去模糊模型"
    doc.core_properties.author = "YOLO ALPR Project"
    add_cover(doc)
    add_toc(doc)
    add_overview(doc)
    add_architecture(doc)
    add_windows(doc)
    add_rk3588(doc)
    add_obb(doc)
    add_deblur(doc)
    add_difficulties(doc)
    add_runbook(doc)
    add_roadmap(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
