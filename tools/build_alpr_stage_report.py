from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "汇报" / "ALPR_板端移植_MIPI_去模糊阶段汇报_20260627.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WARN = "FFF4CE"
OK = "EAF4EA"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_paragraph(doc, text="", style=None, size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    set_table_borders(table, color="D9E2EC", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.3, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    set_table_borders(table, color="D7DBE2", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.strip("\n").splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", east_asia="Microsoft YaHei", size=8.5, color=RGBColor(32, 32, 32))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=font_size, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    return table


def add_image(doc, path, caption, width=6.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        s = styles[style_name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    header = sec.header.paragraphs[0]
    header.text = "ALPR 阶段汇报与演示手册"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.text = "D:/YOLO_ALPR_Project | RK3588 deployment notes"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=8.5, color=MUTED)
    return doc


def summarize_deblur():
    test_metrics = ROOT / "blur models" / "deblur_v2_e20" / "test_metrics.json"
    if test_metrics.exists():
        data = json.loads(test_metrics.read_text(encoding="utf-8"))
    else:
        data = {}
    topk_metrics = ROOT / "blur models" / "deblur_v2_e20" / "eval_captures_topk_all" / "metrics.csv"
    ratios = []
    contrast = []
    if topk_metrics.exists():
        with topk_metrics.open("r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f):
                try:
                    ratios.append(float(row["sharpness_ratio"]))
                    contrast.append(float(row["contrast_ratio"]))
                except Exception:
                    pass
    static_csv = ROOT / "blur models" / "deblur_v2_e20" / "static_hyperlpr3_test" / "results.csv"
    static = {"total": 0, "changed": 0, "orig_ok": 0, "deblur_ok": 0}
    if static_csv.exists():
        with static_csv.open("r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f):
                static["total"] += 1
                static["changed"] += 1 if row.get("changed") == "True" else 0
                static["orig_ok"] += 1 if row.get("original_plate_like") == "True" else 0
                static["deblur_ok"] += 1 if row.get("deblur_plate_like") == "True" else 0
    return data, ratios, contrast, static


def main():
    doc = setup_doc()

    # Cover
    add_paragraph(doc, "技术阶段汇报 / 演示手册", size=11, color=MUTED, bold=True, after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("高速路固定摄像头场景 ALPR：Windows 主线到 RK3588 板端移植、MIPI 实时识别与去模糊实验")
    set_run_font(r, size=23, color=INK, bold=True)
    add_paragraph(doc, "项目目录：D:/YOLO_ALPR_Project", size=10.5, color=MUTED, after=2)
    add_paragraph(doc, f"整理日期：{datetime.now():%Y-%m-%d}", size=10.5, color=MUTED, after=12)
    add_callout(
        doc,
        "一句话结论",
        "Windows Top-K 主线已经迁移到 RK3588 的视频文件识别链路，并形成了速度模式与等价诊断模式两套命令；MIPI 手机图片调试已能直接扫到车牌，但它是临时调试链路，不等同于高速路真实部署；去模糊模型在静态指标和视觉锐度上有效，但对 HyperLPR3 OCR 的收益尚不稳定，暂不建议直接进入板端实时主链路。",
        fill=OK,
    )
    add_table(
        doc,
        ["模块", "当前状态", "演示价值", "主要风险"],
        [
            ["Windows 主线", "Top-K、OCR 投票、锁定后显示逻辑已稳定作为基准", "可展示算法从单帧到多帧投票的演进", "主脚本 alpr_topk_capture.py 不应随意污染"],
            ["RK 视频识别", "可跑完整 144.mp4；i3 + rk_adaptive 历史完整视频约 11.72 FPS", "可展示板端部署结果和浏览器预览", "严格 i1 等价模式约 4.97 FPS，速度明显下降"],
            ["MIPI 实时识别", "颜色已通过 UYVY + gray-world 临时修正；手机图片可 direct plate scan", "可现场展示摄像头画面和实时识别", "手机屏幕里的整车不容易被车辆检测模型识别"],
            ["Deblur", "ONNX 模型已训练并可静态测试；PSNR/SSIM 提升明显", "可展示 before/after 和 OCR A/B", "视觉更锐不必然提升 OCR，板端实时接入成本高"],
        ],
        [1.25, 2.15, 1.55, 1.4],
    )

    doc.add_page_break()
    add_heading(doc, "1. 项目目标与当前边界", 1)
    add_paragraph(doc, "最终目标是在高速路固定摄像头场景下，尽可能接近原视频速度完成中文车牌识别。系统应在车辆进入有效识别区域时完成车牌检测与 OCR，得到足够稳定的投票结果后锁定文本；锁定后不再对该 track 重复做 OBB/OCR，只保留车辆跟踪和车牌号显示。")
    for item in [
        "Windows 端用于算法验证、A/B 试验和去模糊离线评估。",
        "RK3588 端用于最终实时部署，包括视频文件回放和 MIPI 摄像头输入。",
        "MIPI 手机图片识别是调试摄像头和实时流能力的辅助场景，不应直接等同于高速路真实车辆识别。",
        "去模糊当前作为离线复核和后台实验分支，不进入板端实时主链路。",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "当前边界", "现在主线应保持两份脚本：rk3588_topk_capture.py 保留为 Git 提交版/视频部署主线；rk3588_topk_capture_mipi_debug.py 保留 MIPI 手机图片调试直扫能力。这样不会再因为 MIPI 调试污染视频文件识别。", fill=CALLOUT)

    add_heading(doc, "2. 从 Windows 到 RK3588 的移植过程", 1)
    add_heading(doc, "2.1 Windows 主线如何形成", 2)
    for item in [
        "早期 TinyUNet / 图像增强路线对论文数据集有效，但对真实高速路视频 crop 泛化不足。",
        "早期 alpr_sahi_snapshot2.py 可截取车辆和车牌，但单帧 fallback 经常模糊，无法保证 OCR 稳定。",
        "主线转向 Top-K：同一车辆跨多帧收集多个 plate crop，按面积、清晰度、曝光、对比度、OBB 置信度排序。",
        "OCR 从 PaddleOCR、plate-rec 系列逐步转向 HyperLPR3；HyperLPR3 对中文车牌整体最可靠，但仍需要投票抑制噪声。",
        "最终形成 alpr_topk_capture.py：车辆检测、轻量 tracker、OBB、透视矫正、Top-K 保存、OCR 投票、锁定后显示。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 移植到 RK 的关键改造", 2)
    add_table(
        doc,
        ["改造点", "Windows 端", "RK3588 端实现", "遇到的问题"],
        [
            ["推理框架", "Ultralytics / PyTorch / ONNXRuntime", "RKNNLite 加载 vehicle.rknn 与 best_obb.rknn", "RKNN 输出格式、静态 shape 警告、模型版本提示需要确认"],
            ["车辆检测", "YOLO yolo11n.pt", "vehicle.rknn，输入 640", "每帧检测速度较慢，需要 interval 与 adaptive"],
            ["车牌 OBB", "best_obb.pt", "best_obb.rknn / plate_imgsz 640", "命中次数少时 OCR 票数不足"],
            ["OCR", "HyperLPR3", "板端 Python HyperLPR3", "中文输出在部分日志/JSON 中会出现编码显示问题"],
            ["跟踪", "轻量 IoU tracker", "IoU + center/scale + locked reassociation", "断轨会导致锁定文字只闪一下或转移困难"],
            ["实时显示", "OpenCV window", "/tmp/frame.jpg + stream.py", "JPEG 发布会带来小幅 CPU 开销"],
        ],
        [1.05, 1.45, 1.65, 2.2],
    )

    add_heading(doc, "3. 板端视频文件识别现状", 1)
    add_paragraph(doc, "板端现有视频为 /root/deploy/144.mp4，与 Windows 端 D:/YOLO_ALPR_Project/测试图/14.mp4 内容相同。当前应区分两类运行模式：严格等价诊断模式和板端速度/效果平衡模式。")
    add_table(
        doc,
        ["模式", "核心参数", "用途", "观察到的 FPS/现象"],
        [
            ["Windows 等价诊断", "vehicle-detect-interval=1, vehicle_conf=0.45, min_process_vehicle_conf=0.45", "定位差异来源，不追求速度", "1000 帧约 4.97 FPS；低阈值带来更多后续处理"],
            ["RK 速度平衡", "vehicle-detect-interval=3, --rk-adaptive, publish_interval=2", "接近实时演示和完整视频测试", "完整视频 11030 帧平均约 11.72 FPS；能锁定多辆车"],
            ["no-OCR 诊断", "ocr-engine none / lock-on-plate-detect", "证明检测后跟踪显示链路可行", "约 12 FPS，但只能显示 PLATE，不是最终目标"],
            ["MIPI 手机调试", "mipi-plate-scan-only / direct plate scan", "调摄像头和手机图片展示", "约 6-7 FPS；非高速路真实视频主线"],
        ],
        [1.25, 2.2, 1.55, 1.55],
    )
    add_callout(doc, "为什么 4.7 FPS 和 11 FPS 都是“正常”的", "4.7 FPS 来自每帧车辆检测的 Windows 等价诊断模式；11 FPS 来自 vehicle-detect-interval=3 + rk_adaptive 的板端速度模式。二者目标不同，不能直接比较。", fill=WARN)

    add_heading(doc, "3.1 历史 RK 测试结果摘要", 2)
    add_table(
        doc,
        ["测试", "帧数", "FPS", "关键参数", "锁定结果/结论"],
        [
            ["rk_windows_equiv_i1", "1000", "4.97", "interval=1", "可锁定前两辆，但速度慢，适合作为等价诊断"],
            ["rk_adaptive_i3", "1000", "12.72", "interval=3, rk_adaptive", "锁定 冀JC5210、冀B6R9F9，适合板端策略"],
            ["full video run_20260626_074232", "11030", "11.72", "interval=3, rk_adaptive, publish_interval=2", "锁定 冀JC5210、冀B6R9F9、鲁V0JU1Q、冀CH7V97"],
            ["noocr_plate_lock_ab", "1000", "12.03", "no OCR / plate detect lock", "证明锁定后跟踪显示可行，但不能作为最终识别结果"],
        ],
        [1.75, 0.65, 0.65, 1.85, 2.0],
    )

    add_heading(doc, "4. MIPI 摄像头识别现状", 1)
    add_paragraph(doc, "MIPI 调试经历了两个阶段：先解决浏览器是否真的看到实时摄像头画面，再定位为什么手机图片里的车辆难以进入正常车辆检测链路。")
    for item in [
        "stream.py 只是浏览器服务，读取 /tmp/frame.jpg；真正更新画面的程序必须另行运行。",
        "MIPI 采集使用 /dev/video22、UYVY、v4l2-ctl 后端；OpenCV V4L2 路径在板端偏慢。",
        "颜色发绿/发紫主要来自 UYVY 解码、白平衡和曝光问题；已用 robust gray-world 做临时软件白平衡。",
        "手机屏幕中的车辆不等同于真实车辆：车辆检测模型容易失败，但车牌 OBB 可直接在 ROI 中检测到车牌。",
        "因此新增了单独调试版 rk3588_topk_capture_mipi_debug.py，支持 mipi-plate-scan-only，避免污染视频主线。",
    ]:
        add_bullet(doc, item)
    mipi_img = ROOT / "RK3588_dev" / "mipi_current_probe_live" / "annotated_probe.jpg"
    if mipi_img.exists():
        add_image(doc, mipi_img, "图 1：MIPI 当前帧 direct plate scan 诊断。车辆检测为空，但直接车牌 OBB + HyperLPR3 能识别手机画面中的车牌。", width=6.1)
    mipi_fast = ROOT / "RK3588_dev" / "mipi_scan_fast_current.jpg"
    if mipi_fast.exists():
        add_image(doc, mipi_fast, "图 2：MIPI scan-only 轻量配置画面。该模式用于手机图片演示，不代表高速路视频主流程。", width=6.1)

    add_heading(doc, "4.1 MIPI 遇到的困难与尝试", 2)
    add_table(
        doc,
        ["问题", "现象", "尝试", "当前判断"],
        [
            ["浏览器画面不是实时", "stream.py 只服务旧 /tmp/frame.jpg", "写 mipi_preview_writer.py 持续更新帧", "需要区分采集程序和浏览器服务"],
            ["画面发绿/偏暗", "室内背光、手机屏幕反光、白平衡漂移", "UYVY 解码、gray-world、low-light gamma/CLAHE", "可临时改善，但最终应确认 ISP/IQ/曝光"],
            ["车辆检测失败", "手机屏幕里的整车没有 vehicle box", "读取当前帧跑一帧探针", "车检失败但车牌 OBB 可成功"],
            ["旧车牌锁定不释放", "切换手机图片后仍显示上一辆", "MIPI 调试版增加换牌重置", "已隔离在 mipi_debug 脚本中"],
            ["FPS 较低", "全帧 OBB+OCR 约 2 FPS", "scan-only、ROI、降低识别频率", "提升到约 6-7 FPS，仍属调试模式"],
        ],
        [1.25, 1.45, 1.65, 2.1],
    )

    add_heading(doc, "5. 去模糊模型实验", 1)
    data, ratios, contrast, static = summarize_deblur()
    restored = data.get("restored", {})
    baseline = data.get("identity_blur_baseline", {})
    add_paragraph(doc, "去模糊模型 deblur_v2_e20 已完成训练，并导出 ONNX：D:/YOLO_ALPR_Project/blur models/deblur_v2_e20/plate_restore_lite_v2_e20_320x96.onnx。当前策略是先在 Windows 端离线评估，不直接接入 RK 实时主链路。")
    add_table(
        doc,
        ["指标", "输入模糊 baseline", "去模糊 restored", "结论"],
        [
            ["Loss", f"{baseline.get('loss', 0):.4f}", f"{restored.get('loss', 0):.4f}", "损失下降"],
            ["PSNR", f"{baseline.get('psnr', 0):.2f}", f"{restored.get('psnr', 0):.2f}", "约 +4.01 dB"],
            ["SSIM", f"{baseline.get('ssim', 0):.3f}", f"{restored.get('ssim', 0):.3f}", "结构相似度明显提升"],
            ["真实 Top-K crop 锐度", "-", f"平均倍率约 {sum(ratios)/len(ratios):.2f}x" if ratios else "-", "视觉锐度通常提升"],
            ["静态 HyperLPR3 A/B", f"{static['orig_ok']}/{static['total']} plate-like", f"{static['deblur_ok']}/{static['total']} plate-like", "OCR 收益不稳定，需继续分层评估"],
        ],
        [1.45, 1.35, 1.35, 2.25],
    )
    preview = ROOT / "blur models" / "deblur_v2_e20" / "eval_captures_topk_all" / "preview_grid.jpg"
    if preview.exists():
        add_image(doc, preview, "图 3：deblur_v2_e20 对真实 Top-K plate crop 的批量 before/after 预览。", width=5.6)
    add_callout(doc, "去模糊当前结论", "模型在合成测试集和真实 crop 的视觉锐度上有效，但 HyperLPR3 对去模糊后的文本并非总是更稳定；有些样本置信度提升，有些样本会把字符推向另一个错误结果。因此下一步应作为后台复核票源或离线 Top-K 复核层，而不是立即放入板端实时主链路。", fill=WARN)

    add_heading(doc, "6. 演示运行命令", 1)
    add_heading(doc, "6.1 板端视频速度/效果平衡模式", 2)
    add_code(doc, r"""
cd /root/alpr_topk_rk3588

python3 rk3588_topk_capture.py \
  --video /root/deploy/144.mp4 \
  --vehicle-model /root/deploy/vehicle.rknn \
  --plate-model /root/deploy/best_obb.rknn \
  --ocr-engine hyperlpr3 \
  --output /root/alpr_topk_rk3588/runs_rk_adaptive_video \
  --vehicle-detect-interval 3 \
  --rk-adaptive \
  --publish-frame /tmp/frame.jpg \
  --publish-interval 2 \
  --progress-interval 100
""")
    add_heading(doc, "6.2 板端 Windows 等价诊断模式", 2)
    add_code(doc, r"""
cd /root/alpr_topk_rk3588

python3 rk3588_topk_capture.py \
  --video /root/deploy/144.mp4 \
  --vehicle-model /root/deploy/vehicle.rknn \
  --plate-model /root/deploy/best_obb.rknn \
  --ocr-engine hyperlpr3 \
  --output /root/alpr_topk_rk3588/runs_rk_video_equiv \
  --vehicle-detect-interval 1 \
  --vehicle-conf 0.45 \
  --min-process-vehicle-conf 0.45 \
  --min-ocr-conf 0.70 \
  --vote-window 10 \
  --vote-threshold 3 \
  --min-char-vote-ratio 0.65 \
  --publish-frame /tmp/frame.jpg \
  --publish-interval 2 \
  --progress-interval 100
""")
    add_heading(doc, "6.3 浏览器实时流服务", 2)
    add_code(doc, r"""
cd /root/deploy
python3 stream.py

# Windows 浏览器打开：
http://192.168.137.168:8080
""")
    add_heading(doc, "6.4 MIPI 手机图片调试模式", 2)
    add_code(doc, r"""
cd /root/alpr_topk_rk3588

python3 rk3588_topk_capture_mipi_debug.py \
  --video mipi \
  --camera-width 1280 \
  --camera-height 720 \
  --vehicle-model /root/deploy/vehicle.rknn \
  --plate-model /root/deploy/best_obb.rknn \
  --ocr-engine hyperlpr3 \
  --output /root/alpr_topk_rk3588/runs_mipi_phone_scan \
  --mipi-fourcc UYVY \
  --mipi-color-mode uyvy \
  --mipi-backend v4l2ctl \
  --mipi-gray-world \
  --rk-adaptive \
  --mipi-plate-scan-only \
  --mipi-plate-scan-interval 6 \
  --mipi-plate-scan-conf 0.20 \
  --mipi-plate-scan-roi 0.15 0.35 0.90 0.95 \
  --mipi-plate-scan-reset-conf 0.85 \
  --mipi-plate-scan-reset-votes 2 \
  --vote-threshold 2 \
  --adaptive-min-exact-votes 2 \
  --adaptive-min-strong-votes 2 \
  --adaptive-min-vote-frames 2 \
  --min-ocr-conf 0.70 \
  --publish-frame /tmp/frame.jpg \
  --publish-interval 1 \
  --preview-jpeg-quality 78 \
  --progress-interval 30
""")
    add_heading(doc, "6.5 Windows 端 deblur 静态图片 + HyperLPR3 测试", 2)
    add_code(doc, r"""
cd D:\YOLO_ALPR_Project

D:\miniconda\envs\alpr_env\python.exe .\test_deblur_hyperlpr3.py `
  --input "D:\YOLO_ALPR_Project\RK3588_dev\run_20260626_074232" `
  --recursive `
  --output "D:\YOLO_ALPR_Project\blur models\deblur_v2_e20\static_hyperlpr3_test"

# 结果查看：
# restored\        去模糊后单图
# comparison\      原图/去模糊对照图
# results.csv      原图 OCR 与去模糊 OCR 对比
# results.json     结构化结果
""")
    add_heading(doc, "6.6 Windows demo 主流程接入 deblur 后台投票实验", 2)
    add_code(doc, r"""
cd D:\YOLO_ALPR_Project

D:\miniconda\envs\alpr_env\python.exe .\alpr_topk_capture_demo.py `
  --video "D:\YOLO_ALPR_Project\测试图\14.mp4" `
  --output "D:\YOLO_ALPR_Project\captures_deblur_bg_vote_full" `
  --live-ocr `
  --ocr-engine hyperlpr3 `
  --deblur-model "D:\YOLO_ALPR_Project\blur models\deblur_v2_e20\plate_restore_lite_v2_e20_320x96.onnx" `
  --deblur-mode always `
  --progress-interval 100
""")

    add_heading(doc, "7. 汇报时建议讲法", 1)
    add_number(doc, "先讲目标：不是单帧识别，而是固定摄像头视频中“识别一次、锁定后跟踪显示”。")
    add_number(doc, "再讲 Windows 主线：从单帧 fallback 走向 Top-K、多帧投票、锁定策略。")
    add_number(doc, "展示 RK 视频：先跑速度模式，让听众看到完整链路已经能在板端运行；再解释等价模式为什么慢。")
    add_number(doc, "展示 MIPI：强调这是摄像头实时链路和手机图片调试，不等同于高速路模型效果；说明 direct plate scan 是为了排查车检失败。")
    add_number(doc, "展示 deblur：先展示视觉 before/after，再诚实说明 OCR 收益不稳定，所以暂时放在后台复核/离线层。")
    add_number(doc, "最后讲下一步：OBB 模型上板 A/B、ROI 与调度优化、MIPI ISP 正规修复、deblur 只在有收益时进入后台票源。")

    add_heading(doc, "8. 下一步计划", 1)
    add_table(
        doc,
        ["优先级", "任务", "具体动作", "验收标准"],
        [
            ["P0", "稳定板端视频主线", "保留 Git 版 rk3588_topk_capture.py；每次改动跑 1000 帧以上并保存 run_summary", "速度模式稳定 10-12 FPS，锁牌结果不退化"],
            ["P0", "命令和脚本分层", "主线视频脚本与 MIPI 调试脚本分离；避免调 MIPI 污染视频主线", "演示命令可重复运行，行为可解释"],
            ["P1", "OBB e20 上板 A/B", "将 best_obb_finetuned_e20.pt 转 ONNX/RKNN，与 best_obb.rknn 同视频对比", "plate_hits、rejected、锁定帧号、FPS 有结构化对比"],
            ["P1", "MIPI 正规修复", "确认 media-ctl、v4l2-ctl、rkaiq_3A_server、IQ 文件与 pixel format", "不依赖强软件调色也能获得稳定颜色"],
            ["P1", "deblur 后台复核", "只对未识别或低置信 Top-K crop 做异步 deblur+OCR，再进入二级投票", "证明在真实 Top-K crop 上提升锁定率且 FPS 可接受"],
            ["P2", "跟踪增强", "必要时评估 ByteTrack/OC-SORT 或固定摄像头运动模型", "减少断轨和锁定文本丢失"],
        ],
        [0.65, 1.25, 2.65, 1.8],
    )

    add_heading(doc, "9. 风险与注意事项", 1)
    for item in [
        "不要把 MIPI 手机调试模式当作高速路真实车辆模型效果，它主要验证摄像头、流服务、颜色和车牌直扫能力。",
        "不要随意改 alpr_topk_capture.py；Windows 新功能优先放 alpr_topk_capture_demo.py。",
        "严格等价模式慢是预期现象；演示实时效果应使用 interval=3 + rk_adaptive。",
        "deblur 视觉更锐不等于 OCR 更准；必须用 results.csv 的 OCR A/B 结果说话。",
        "板端测试不要只看前几百帧，至少跑到 frame=1000，完整视频结果更有说服力。",
        "远程 SSH/SCP 涉及网络访问；Codex 执行时需要提升权限。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "附录 A：环境与路径", 1)
    add_table(
        doc,
        ["对象", "路径 / 地址", "说明"],
        [
            ["RK3588 SSH", "root@192.168.137.168", "板端运行和测试"],
            ["Ubuntu 开发机", "apple@192.168.217.128", "Ubuntu 22.04.5 LTS"],
            ["板端脚本目录", "/root/alpr_topk_rk3588", "RK 脚本、runs 输出"],
            ["板端部署目录", "/root/deploy", "视频、模型、stream.py"],
            ["板端测试视频", "/root/deploy/144.mp4", "与 Windows 14.mp4 内容相同"],
            ["Windows 视频", "D:/YOLO_ALPR_Project/测试图/14.mp4", "Windows demo 输入"],
            ["实时浏览器", "http://192.168.137.168:8080", "stream.py 读取 /tmp/frame.jpg"],
            ["deblur 模型", "D:/YOLO_ALPR_Project/blur models/deblur_v2_e20", "ONNX、对照图、metrics"],
        ],
        [1.25, 2.6, 2.15],
    )

    add_heading(doc, "附录 B：当前文件状态建议", 1)
    add_paragraph(doc, "截至本报告整理时，建议保持以下分工：")
    add_bullet(doc, "rk3588_topk_capture.py：保留 Git 提交版，作为板端视频文件识别主线。")
    add_bullet(doc, "rk3588_topk_capture_mipi_debug.py：保留 MIPI 手机图片识别、直扫 ROI、换牌重置等调试能力。")
    add_bullet(doc, "alpr_topk_capture.py：Windows 基准主线，原则上不污染。")
    add_bullet(doc, "alpr_topk_capture_demo.py：Windows 实验分支，可用于 deblur 后台投票实验。")
    add_bullet(doc, "test_deblur_hyperlpr3.py：静态图片去模糊 + HyperLPR3 OCR 对比工具。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
